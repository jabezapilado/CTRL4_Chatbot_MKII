"""
Retrieval-Augmented Generation (RAG) Service

Builds and manages the Guidance Office
knowledge retrieval system.

Responsibilities

- Document Processing
- Text Chunking
- Embedding Generation
- FAISS Index Management
- Semantic Retrieval

CTRL4 Chatbot MK2

Authors:
- Apilado, Jabez Timothy E.
- Quilantang, Grant Mihkael D.
- Lanix, Iligan
- Wylengco, Teyshaun Zell
"""

from __future__ import annotations

import json
import re

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from ..config import Config

@dataclass
class RetrievedDocument:
    text: str
    source: str
    score: float

def normalize_text(text: str) -> str:
    lowered = str(text).strip().lower()
    lowered = re.sub(r"\s+", " ", lowered)
    return lowered

def split_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end]
        if end < len(text):
            split_at = max(chunk.rfind("\n\n"), chunk.rfind(". "), chunk.rfind("? "), chunk.rfind("! "))
            if split_at > chunk_size // 3:
                end = start + split_at + 1
                chunk = text[start:end]
        chunks.append(chunk.strip())
        if end >= len(text):
            break
        start = max(0, end - overlap)

    return [c for c in chunks if c]

def _extract_pdf_text(file_path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        return ""

    pages: list[str] = []
    reader = PdfReader(str(file_path))
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx_text(file_path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return ""

    document = Document(str(file_path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)

def read_document_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(file_path)
    if suffix == ".docx":
        return _extract_docx_text(file_path)
    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="ignore")
    return ""

class RAGService:
    
    def __init__(self):

        self.config = Config()

        self.docs_dir = Path(self.config.RAG_DOCS_DIR)

        self.index_dir = Path(self.config.RAG_INDEX_DIR)

        self.index_file = self.index_dir / "knowledge.faiss"

        self.metadata_file = self.index_dir / "metadata.json"

        self.embedding_model_name = self.config.RAG_EMBEDDING_MODEL

        self.embedder: Any | None = None
        self.faiss: Any | None = None
        self.index: Any | None = None
        self.metadata: list[dict[str, Any]] = []

        self.initialization_error: str | None = None

        self._initialize_runtime()
    
    def _build_embedding_input(self, text: str, is_query: bool) -> str:
        if "e5" in self.embedding_model_name.lower():
            return f"query: {text}" if is_query else f"passage: {text}"
        return text

    def _embed(
        self,
        texts: list[str],
        is_query: bool,
    ) -> Any:

        if self.embedder is None:
            raise RuntimeError(
                "Embedding model is not available."
            )

        prepared = [
            self._build_embedding_input(text, is_query=is_query)
            for text in texts
        ]

        return self.embedder.encode(
            prepared,
            normalize_embeddings=True,
            convert_to_numpy=True,
        )

    def _initialize_runtime(self) -> None:
        try:
            import faiss  # type: ignore
            from sentence_transformers import SentenceTransformer  # type: ignore
        except Exception:
            self.initialization_error = (
                "RAG dependencies are missing. Install sentence-transformers and faiss-cpu."
            )
            return

        self.faiss = faiss
        self.embedder = SentenceTransformer(self.embedding_model_name)

        if self.index_file.exists() and self.metadata_file.exists():
            self._load_index()
            return

        if self.config.RAG_AUTO_BUILD_ON_START:
            self.build_index()

    def _load_index(self) -> None:
        if not self.faiss:
            return
        self.index = self.faiss.read_index(str(self.index_file))
        self.metadata: list[dict[str, Any]] = json.loads(
            self.metadata_file.read_text(encoding="utf-8")
        )

    def build_index(self) -> int:
        if not self.embedder or not self.faiss:
            return 0

        self.index_dir.mkdir(parents=True, exist_ok=True)
        self.docs_dir.mkdir(parents=True, exist_ok=True)

        doc_paths = sorted(
            [
                path
                for path in self.docs_dir.rglob("*")
                if path.is_file() and path.suffix.lower() in {".pdf", ".docx", ".txt", ".md"}
            ]
        )

        chunks: list[dict[str, Any]] = []
        for path in doc_paths:
            raw = read_document_text(path)
            clean = normalize_text(raw)
            if not clean:
                continue
            for idx, chunk in enumerate(
                split_text(
                    clean,
                    chunk_size=self.config.RAG_CHUNK_SIZE,
                    overlap=self.config.RAG_CHUNK_OVERLAP,
                )
            ):
                chunks.append(
                    {
                        "text": chunk,
                        "source": str(path.relative_to(self.docs_dir)),
                        "chunk_id": idx,
                    }
                )

        if not chunks:
            self.index = None
            self.metadata = []
            return 0

        vectors = self._embed([item["text"] for item in chunks], is_query=False)
        dim = int(vectors.shape[1])
        index = self.faiss.IndexFlatIP(dim)
        index.add(vectors)

        self.faiss.write_index(index, str(self.index_file))
        self.metadata_file.write_text(json.dumps(chunks, ensure_ascii=False, indent=2), encoding="utf-8")

        self.index = index
        self.metadata = chunks
        return len(chunks)

    def retrieve(self, query: str) -> list[RetrievedDocument]:

        if not self.embedder or not self.faiss or self.index is None or not self.metadata:
            return []

        query_vector = self._embed([query], is_query=True)

        k = max(1, self.config.RAG_TOP_K)

        scores, indices = self.index.search(query_vector, k)

        results: list[RetrievedDocument] = []

        for score, idx in zip(scores[0], indices[0]):

            if idx < 0:
                continue

            if float(score) < self.config.RAG_MIN_SCORE:
                continue

            meta = self.metadata[int(idx)]

            results.append(
                RetrievedDocument(
                    text=str(meta.get("text", "")),
                    source=str(meta.get("source", "unknown")),
                    score=float(score),
                )
            )

        results.sort(
            key=lambda document: document.score,
            reverse=True
        )

        return results
    
    def reload_index(self) -> int:

        if not self.index_file.exists():
            return 0

        self._load_index()

        return len(self.metadata)
    
    def stats(self) -> dict[str, object]:

        return {
            "documents": len(self.metadata),
            "embedding_model": self.embedding_model_name,
            "indexed": self.index is not None,
        }
        
    @property
    def ready(self) -> bool:
        return (
            self.embedder is not None
            and self.index is not None
            and self.faiss is not None
        )      
    